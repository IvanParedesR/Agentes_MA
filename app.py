import streamlit as st
import pdfplumber
import re
import json
import os
import tempfile
from openai import OpenAI
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
import warnings
import pandas as pd
import tempfile
from io import BytesIO
from datetime import datetime
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os
import easyocr
import fitz
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from openai import OpenAI
from langchain.prompts import PromptTemplate

st.set_page_config(page_title="Merger and Acquisition - Revisión de documentos", layout="wide", page_icon="📄")
# Configuración para Windows (coloca esto ANTES de usar pytesseract)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Verifica la instalación (opcional)
try:
    print(pytesseract.get_tesseract_version())
except:
    st.warning("Tesseract no está instalado correctamente")

# Configuración inicial
warnings.filterwarnings("ignore", category=UserWarning, message=".*CropBox.*")
# Sidebar con configuración
with st.sidebar:
    st.title("APIs de DeepSeek y OpenAI")   
    st.markdown("---")
    api_key_deepseek = st.text_input("API Key DeepSeek", type="password")
    api_key_openai = st.text_input("API Key OpenAI", type="password")
    st.markdown("---")
    st.markdown("### Merger and Adquisicion Analysis:")
    st.markdown("- Pago de derechos")
    st.markdown("- Fusiones y Adquisiciones")
    st.markdown("- Poderes Legales")
    st.markdown("---")
    st.markdown("Tarea de Agentes AI - ITAM 2025 - Iván PR")

# Iniciar API
def get_clients():
    return (
        OpenAI(api_key=api_key_deepseek, base_url="https://api.deepseek.com"),
        OpenAI(api_key=api_key_openai)
    )

# Funciones comunes
reader = easyocr.Reader(['es', 'en'], gpu=False)

def load_document(file_path, file_type):
    text = ""

    if file_type == "application/pdf":
        try:
            # Intentar extracción directa con pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""

            # Si el texto extraído es escaso, usar OCR con EasyOCR
            if len(text.strip()) < 50:
                try:
                    images = convert_from_path(file_path)
                    for image in images:
                        result = reader.readtext(np.array(image), detail=0)
                        text += "\n".join(result) + "\n"
                except Exception as ocr_error:
                    st.warning(f"Ocurrió un error con OCR: {str(ocr_error)}")

        except Exception as e:
            st.error(f"Error al leer PDF: {str(e)}")

    elif file_type.startswith("image/"):
        try:
            image = Image.open(file_path)
            result = reader.readtext(np.array(image), detail=0)
            text = "\n".join(result)
        except Exception as img_error:
            st.error(f"No se pudo procesar la imagen: {str(img_error)}")

    else:  # Archivo de texto plano
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            st.error(f"Error al leer archivo de texto: {str(e)}")

    return text

# Funciones para cada tipo de análisis
def analisis_depositos(uploaded_file):
    client, _ = get_clients()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        file_path = tmp_file.name
    
    document_text = load_document(file_path, uploaded_file.type)
    
    with st.expander("🔍 Ver texto extraído (primeros 1000 caracteres)"):
        st.text(document_text[:1000] + ("..." if len(document_text) > 1000 else ""))
    
    if st.button("🔎 Analizar Depósito"):
        with st.spinner("🧠 Analizando con DeepSeek..."):
            prompt = """Extrae EXACTAMENTE en formato JSON (sin texto adicional) estos datos:
            {
                "beneficiario": "Nombre completo",
                "ordenante": "Nombre completo",
                "monto": "Cantidad con moneda (ej. $1,000.00 MXN)",
                "fecha_hora": "DD/MM/AAAA HH:MM"
            }
            Texto del comprobante:
            """ + document_text[:4000]  # Limitar tamaño para no exceder tokens
            
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system", 
                        "content": "Extrae solo datos de depósitos bancarios. Devuelve ÚNICAMENTE el JSON válido, sin comentarios."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,  # Menor temperatura para más precisión
                response_format={"type": "json_object"},  # Fuerza formato JSON
                max_tokens=500
            )
            
            try:
                # Inspeccionar respuesta cruda primero
                raw_response = response.choices[0].message.content
                st.expander("Respuesta cruda del modelo").code(raw_response)
                
                # Parsear JSON
                data = json.loads(raw_response)
                
                # Validar campos
                required_fields = ["beneficiario", "ordenante", "monto", "fecha_hora"]
                for field in required_fields:
                    if field not in data:
                        st.warning(f"Campo faltante: {field}")
                        data[field] = "No especificado"
                    elif not data[field].strip():
                        data[field] = "No especificado"
                
                # Mostrar resultados
                st.success("✅ Información del depósito")
                cols = st.columns(2)
                with cols[0]:
                    st.metric("🧑 Beneficiario", data["beneficiario"])
                    st.metric("👤 Ordenante", data["ordenante"])
                with cols[1]:
                    st.metric("💰 Monto", data["monto"])
                    st.metric("📅 Fecha y Hora", data["fecha_hora"])
                
                # Descarga
                st.download_button(
                    label="⬇️ Descargar JSON",
                    data=json.dumps(data, indent=2, ensure_ascii=False),
                    file_name="deposito.json",
                    mime="application/json"
                )
                
            except json.JSONDecodeError as e:
                st.error(f"❌ El modelo no devolvió un JSON válido: {e}")
                st.code(f"Respuesta problema:\n{raw_response}")
            except Exception as e:
                st.error(f"Error inesperado: {e}")
                st.stop()

def generar_dictamen_formateado(data, narrativa):
    try:
        doc = Document()
        
        # Estilo general
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Título
        title = doc.add_heading('DICTAMEN SIMPLIFICADO', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style.font.bold = True
        
        # Encabezado con formato
        header = doc.add_paragraph()
        header.add_run(f"Expediente: CNT-{datetime.now().strftime('%Y')}-XXXX\n")
        header.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}\n")
        header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Sección I. Antecedentes
        doc.add_heading('I. Antecedentes', level=2)
        for item in data.get('antecedentes', []):
            doc.add_paragraph(item, style='List Number')
        
        # Sección II. Consideraciones de Derecho
        doc.add_heading('II. Consideraciones de Derecho', level=2)
        for item in data.get('consideraciones', []):
            doc.add_paragraph(item, style='List Number')
        
        # Resolución
        doc.add_heading('RESUELVE', level=2)
        for item in data.get('resoluciones', []):
            doc.add_paragraph(item, style='List Number')
        
        # Tabla de partes (ejemplo)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Agente'
        hdr_cells[1].text = 'País'
        hdr_cells[2].text = 'Actividades'
        hdr_cells[3].text = 'Participación'
        
        # Añadir datos a la tabla
        for parte in data.get('partes', []):
            row_cells = table.add_row().cells
            row_cells[0].text = parte.get('nombre', '')
            row_cells[1].text = parte.get('pais', '')
            row_cells[2].text = parte.get('actividades', '')
            row_cells[3].text = parte.get('participacion', '')
        
        # Notas al pie
        doc.add_paragraph("\nNotas:", style='Intense Quote')
        for nota in data.get('notas', []):
            doc.add_paragraph(nota, style='List Bullet')

        # Opción 1: Guardar en BytesIO (recomendado para Streamlit)
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Botón de descarga
        st.download_button(
            label="⬇️ Descargar Dictamen Completo",
            data=doc_bytes,
            file_name=f"Dictamen_CNT_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        return True

    except Exception as e:
        st.error(f"Error al generar el documento: {str(e)}")
        st.error(f"Detalle técnico: {type(e).__name__} en línea {e.__traceback__.tb_lineno}")
        return False


def analisis_fusiones(uploaded_file):
    client, client_oa = get_clients()

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        file_path = tmp_file.name

    document_text = load_document(file_path, uploaded_file.type)

    if st.button("Analizar Fusión"):
        with st.spinner("🧠 Extrayendo información..."):
            prompt = """Analiza el texto y extrae información en JSON con:
            - "compradores": lista de compradores
            - "vendedores": lista de vendedores
            - "operacion": descripción
            - "monto": valor con moneda
            - "traslapes_horizontales": existencia y descripción
            - "integraciones_verticales": existencia y descripción
            - "apoderados_legales": lista de apoderados"""

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "Eres un experto en fusiones y adquisiciones."},
                    {"role": "user", "content": f"{prompt}\n\n{document_text}"}
                ],
                temperature=0.3,
                max_tokens=1000
            )

            try:
                json_str = re.search(r'\{.*\}', response.choices[0].message.content, re.DOTALL).group()
                data = json.loads(json_str)

                st.success("Información de la operación:")
                st.json(data)

                # Generar narrativa
                with st.spinner("Generando informe ejecutivo..."):
                    narrativa = client_oa.chat.completions.create(
                        model="gpt-4-turbo",
                        messages=[
                            {"role": "system", "content": "Eres un analista financiero experto."},
                            {"role": "user", "content": f"Crea un informe ejecutivo basado en:\n{json.dumps(data, indent=2)}"}
                        ]
                    ).choices[0].message.content

                    st.subheader("Informe Ejecutivo")
                    st.write(narrativa)

                    # Generar documento Word con formato específico
                    if st.button("📄 Generar Dictamen Formateado"):
                        output_path = r"C:\Users\CFC\Downloads\Dictamen_CNT_{}.docx".format(
                            datetime.now().strftime("%Y%m%d_%H%M%S"))
                        
                        # Preparar datos para el formato específico
                        formatted_data = {
                            "antecedentes": [
                                f"El {datetime.now().strftime('%d de %B de %Y')}, {', '.join(data.get('compradores', []))} notificó a esta Comisión su intención de realizar una concentración.",
                                f"La operación consiste en: {data.get('operacion', '')}"
                            ],
                            "consideraciones": [
                                "La Comisión tiene a su cargo la prevención de concentraciones cuyo objeto o efecto sea obstaculizar la competencia económica.",
                                f"La operación notificada consiste en: {data.get('operacion', '')}"
                            ],
                            "resoluciones": [
                                "Autorizar la realización de la concentración notificada.",
                                "Establecer un plazo de vigencia de seis meses para la autorización."
                            ],
                            "partes": [
                                {
                                    "nombre": ", ".join(data.get('compradores', [])),
                                    "pais": "México",
                                    "actividades": "Varios sectores industriales",
                                    "participacion": "100%"
                                },
                                {
                                    "nombre": ", ".join(data.get('vendedores', [])),
                                    "pais": "México",
                                    "actividades": "Varios sectores industriales",
                                    "participacion": "100%"
                                }
                            ],
                            "notas": [
                                "Este dictamen se emite en cumplimiento del artículo 86 de la LFCE.",
                                f"Monto de la operación: {data.get('monto', 'No especificado')}"
                            ]
                        }
                        
                        generar_dictamen_formateado(formatted_data, narrativa, output_path)
                        st.success(f"Dictamen generado y guardado en: {output_path}")

            except Exception as e:
                st.error(f"Error al procesar: {e}")
            finally:
                try:
                    os.unlink(file_path)
                except:
                    pass

import easyocr
import tempfile
import json
import streamlit as st
from funciones import load_document, get_clients  # Asegúrate de tener load_document actualizado

# Inicializa el lector de EasyOCR (fuera de funciones para no reinicializar)
reader = easyocr.Reader(['es', 'en'], gpu=False)

def analisis_poderes(uploaded_file):
    client, _ = get_clients()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        file_path = tmp_file.name

    # Usa OCR con EasyOCR si es necesario
    document_text = load_document(file_path, uploaded_file.type)

    # Mostrar texto extraído
    with st.expander("Ver texto extraído (primeros 500 caracteres)"):
        st.text(document_text[:500] + ("..." if len(document_text) > 500 else ""))

    if st.button("Analizar Poder Legal"):
        with st.spinner("Extrayendo información del poder..."):
            prompt = """Analiza este texto legal y extrae en formato JSON:
            {
                "otorgante": "nombre completo",
                "apoderado": "nombre completo",
                "fecha_inicio": "DD/MM/AAAA",
                "facultades": ["lista", "de", "facultades"],
                "facultades_administracion": ["lista"],
                "limitaciones": "texto descriptivo",
                "observaciones": "problemas detectados"
            }
            Texto:""" + document_text[:13000]

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": "Extrae información legal precisa. Devuelve SOLO el JSON."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                response_format={"type": "json_object"},
                max_tokens=1000
            )

            try:
                raw_response = response.choices[0].message.content
                data = json.loads(raw_response)

                # Validación de campos clave
                for field in ["otorgante", "apoderado", "fecha_inicio"]:
                    if field not in data:
                        data[field] = "No especificado"

                # Mostrar resultados
                st.success("✅ Información del poder legal")
                st.metric("📅 Fecha de inicio", data["fecha_inicio"])

                st.subheader("👤 Otorgante")
                st.write(data.get("otorgante", []))

                st.subheader("👤 Apoderado")
                st.write(data.get("apoderado", []))

                st.subheader("📋 Facultades Principales")
                st.write(data.get("facultades", []))

                st.subheader("🏛️ Facultades de Administración")
                st.write(data.get("facultades_administracion", []))

                if data.get("limitaciones"):
                    st.subheader("🚫 Limitaciones")
                    st.write(data["limitaciones"])

                if data.get("observaciones"):
                    st.warning("⚠️ Observaciones")
                    st.write(data["observaciones"])

            except json.JSONDecodeError:
                st.error("El modelo no devolvió un JSON válido")
                st.code(f"Respuesta cruda:\n{raw_response}")
            except Exception as e:
                st.error(f"Error inesperado: {str(e)}")
# Agregar esta función con las demás funciones de análisis
def analisis_rag(uploaded_file):
    client, _ = get_clients()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        file_path = tmp_file.name
    
    # Extraer texto del PDF
    def extraer_texto_pdf(ruta_pdf):
        doc = fitz.open(ruta_pdf)
        texto = ""
        for pagina in doc:
            texto += pagina.get_text()
        return texto
    
    # Fragmentar texto en chunks
    def fragmentar_texto(texto):
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )
        documentos = [Document(page_content=chunk) for chunk in splitter.split_text(texto)]
        return documentos
    
    # Crear vectorstore FAISS
    def crear_vectorstore(documentos):
        embeddings = OpenAIEmbeddings(openai_api_key=api_key_openai)
        vectorstore = FAISS.from_documents(documentos, embeddings)
        return vectorstore
    
    # Crear cadena QA RAG
    def crear_cadena_qa(vectorstore):
        llm = ChatOpenAI(model="gpt-4", temperature=0, openai_api_key=api_key_openai)

        prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""
            Eres un experto en Fusiones y Adquisiciones con amplia experiencia analizando contratos. 
            Utiliza el siguiente contexto para responder la pregunta legal de forma precisa, técnica y clara.
            Siempre referencia el número de cláusula o artículo donde se encuentra la información.

            Contexto:
            {context}

            Pregunta:
            {question}

            Respuesta como experto:
            """
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=vectorstore.as_retriever(),
            chain_type="stuff",
            chain_type_kwargs={"prompt": prompt_template}
        )

        return qa_chain
    
    # Preguntas predefinidas
    preguntas = [
        "¿Quiénes son los firmantes en el contrato? Describelas y explica su papel. Señala el número de las cláusulas a las que haces referencia.",
        "¿Qué operación, fusión o merger es la que se está acordando o realizando? Señala el número de las cláusulas a las que haces referencia.",
        "¿Hay cláusulas de no competencia? Señala el número de las cláusulas a las que haces referencia.",
        "¿Hay cláusulas de no solicitación? Señala el número de las cláusulas a las que haces referencia.",
        "¿Se requiere aprobación regulatoria para la realización de la fusión u operación? Señala el número de las cláusulas a las que haces referencia."
    ]
    
    if st.button("🔍 Analizar con RAG"):
        with st.spinner("Procesando documento con RAG..."):
            try:
                # Procesar el documento
                texto = extraer_texto_pdf(file_path)
                documentos = fragmentar_texto(texto)
                vectorstore = crear_vectorstore(documentos)
                qa_chain = crear_cadena_qa(vectorstore)
                
                # Mostrar preguntas y respuestas
                st.subheader("Análisis RAG del Contrato")
                
                for pregunta in preguntas:
                    with st.expander(f"❓ {pregunta}"):
                        respuesta = qa_chain.run(pregunta)
                        st.write(respuesta)
                
                # Opción para preguntas personalizadas
                st.markdown("---")
                st.subheader("Realizar pregunta personalizada")
                pregunta_personalizada = st.text_input("Escribe tu pregunta sobre el contrato:")
                
                if pregunta_personalizada:
                    with st.spinner("Buscando respuesta..."):
                        respuesta_personalizada = qa_chain.run(pregunta_personalizada)
                        st.success("Respuesta:")
                        st.write(respuesta_personalizada)
                        
            except Exception as e:
                st.error(f"Error en el análisis RAG: {str(e)}")
            finally:
                try:
                    os.unlink(file_path)
                except:
                    pass
# Interfaz principal
st.title("📄 M&A Analysis")
st.markdown("Analiza documentos financieros, legales y corporativos con IA")

# Inicializar session state para los uploaders
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = {
        'deposito': None,
        'fusion': None,
        'poder': None,
        'rag': None
    }
# Agregar la nueva pestaña RAG
tab1, tab2, tab3, tab4 = st.tabs([
    "Depósitos Bancarios", 
    "Fusiones y Adquisiciones", 
    "Poderes Legales",
    "Análisis de Contratos"
])

with tab1:
    st.header("📊 Análisis de Depósitos Bancarios")
    uploaded_deposito = st.file_uploader("Sube un comprobante de depósito", 
                                       type=["pdf", "png", "jpg", "txt"], 
                                       key="deposito_uploader")
    if uploaded_deposito:
        st.session_state.uploaded_files['deposito'] = uploaded_deposito

with tab2:
    st.header("📜 Análisis de Fusiones y Adquisiciones")
    uploaded_fusion = st.file_uploader("Sube un documento de M&A", 
                                     type=["pdf", "txt"], 
                                     key="fusion_uploader")
    if uploaded_fusion:
        st.session_state.uploaded_files['fusion'] = uploaded_fusion

with tab3:
    st.header("⚖️ Análisis de Poderes Legales")
    uploaded_poder = st.file_uploader("Sube un poder legal", 
                                    type=["pdf", "png", "jpg", "txt"], 
                                    key="poder_uploader")
    if uploaded_poder:
        st.session_state.uploaded_files['poder'] = uploaded_poder
with tab4:
    st.header("🧠 Análisis del Contrato de la Concentración")
    st.markdown("""
    **Análisis profundo de contratos de fusión usando con (RAG)**
    - Identificación precisa de partes involucradas
    - Detección de cláusulas clave (no competencia, no solicitación)
    - Análisis de condiciones regulatorias
    """)
    
    uploaded_rag = st.file_uploader("Sube un contrato de fusión para análisis avanzado", 
                                  type=["pdf"], 
                                  key="rag_uploader")
    if uploaded_rag:
        st.session_state.uploaded_files['rag'] = uploaded_rag

# Procesar archivos subidos
if st.session_state.uploaded_files['deposito'] is not None:
    analisis_depositos(st.session_state.uploaded_files['deposito'])
if st.session_state.uploaded_files['fusion'] is not None:
    analisis_fusiones(st.session_state.uploaded_files['fusion'])
if st.session_state.uploaded_files['poder'] is not None:
    analisis_poderes(st.session_state.uploaded_files['poder'])
if st.session_state.uploaded_files['rag'] is not None:
    analisis_rag(st.session_state.uploaded_files['rag'])
# Estilos CSS personalizados
st.markdown("""
    <style>
    .stMetric {
        background-color: #f0f2f6;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 15px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    .stMetric label {
        font-size: 14px !important;
        color: #555 !important;
    }
    .stMetric div {
        font-size: 18px !important;
        font-weight: bold !important;
        color: #222 !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 10px 20px;
        border-radius: 8px 8px 0 0 !important;
    }
    .stTabs [aria-selected="true"] {
        background-color: #f0f2f6;
    }
    </style>
""", unsafe_allow_html=True)
